#!/usr/bin/env python3
"""Shared helpers for generating AESG CV templates from extracted JSON."""

from __future__ import annotations

import argparse
import base64
import copy
import json
import math
import re
import tempfile
import zipfile
from pathlib import Path, PurePosixPath
from typing import Any, Callable, Dict, Iterable, List, Optional, Tuple

from PIL import Image, ImageOps
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from text_normalization import repair_text_artifacts

AESG_TEAL = "008C95"
TEXT_DARK = "303030"
MUTED = "666666"

STANDARD_SECTION_KEYS = {
    "overview",
    "work experience",
    "professional bodies",
    "professional bodies and memberships",
    "professional memberships",
    "memberships",
    "qualifications",
    "education",
    "core competencies",
    "selected projects",
    "selected project highlights",
    "additional project experience",
}


def load_cvs(path: Path) -> List[Dict[str, Any]]:
    with path.open(encoding="utf-8") as handle:
        payload = json.load(handle)
    if "cvs" in payload:
        return payload["cvs"]
    return [payload]


def load_template_map(path: Optional[Path]) -> Dict[str, Any]:
    if not path:
        return {}
    with path.open(encoding="utf-8") as handle:
        return json.load(handle)


def slugify(value: str) -> str:
    value = re.sub(r"[^A-Za-z0-9]+", "_", value).strip("_")
    return value or "cv"


def source_output_stem(cv: Dict[str, Any]) -> str:
    """Return the source PDF basename exactly, without its extension.

    V3 outputs are already separated by format and orientation folders, so
    retaining the source basename is both unambiguous and easier to trace back
    to the input CV. Older JSON without source metadata keeps a safe slugified
    display-name fallback.
    """
    source = cv.get("source") or {}
    filename = str(source.get("filename") or "").strip()
    if not filename and source.get("file"):
        filename = Path(str(source["file"])).name
    if filename:
        stem = Path(filename).stem
        if stem and stem not in {".", ".."}:
            return stem
    fallback = cv.get("name") or cv.get("id") or "cv"
    return slugify(str(fallback))


def source_output_relative_path(cv: Dict[str, Any], extension: str) -> Path:
    """Return a safe mirrored output path with the requested extension."""
    if not extension.startswith("."):
        extension = f".{extension}"
    source = cv.get("source") or {}
    relative_value = str(source.get("relative_path") or "").strip()
    if relative_value:
        relative = PurePosixPath(relative_value)
        if (
            not relative.is_absolute()
            and relative.name
            and all(part not in {"", ".", ".."} for part in relative.parts)
        ):
            parent = Path(*relative.parent.parts) if relative.parent.parts else Path()
            return parent / f"{relative.stem}{extension}"
    return Path(f"{source_output_stem(cv)}{extension}")


def artifact_output_path(
    output_root: Path,
    cv: Dict[str, Any],
    artifact_folder: str,
    extension: str,
) -> Path:
    """Place an artifact folder inside the mirrored source directory."""
    relative_output = source_output_relative_path(cv, extension)
    return output_root / relative_output.parent / artifact_folder / relative_output.name


def get_section(cv: Dict[str, Any], *headings: str) -> Optional[Dict[str, Any]]:
    wanted = {heading.lower() for heading in headings}
    for section in cv.get("sections", []):
        if section.get("heading", "").lower() in wanted:
            return section
    return None


def section_lines(cv: Dict[str, Any], *headings: str) -> List[str]:
    section = get_section(cv, *headings)
    if not section:
        return []
    lines = [line.get("text", "") for line in section.get("lines", [])]
    return [line for line in lines if line]


def section_text(cv: Dict[str, Any], *headings: str) -> str:
    section = get_section(cv, *headings)
    return (section or {}).get("text", "").strip()


def section_bullets(cv: Dict[str, Any], *headings: str) -> List[str]:
    section = get_section(cv, *headings)
    if not section:
        return []
    bullets = section.get("bullets") or []
    if bullets:
        return bullets
    return section_lines(cv, *headings)


def clean_display_line(value: str) -> str:
    value = repair_text_artifacts(value)
    value = re.sub(r"\s+", " ", value.strip())
    value = re.sub(r"^(?:[•\uf0b7*]|[-–])\s+", "", value).strip()
    return value


def profile(cv: Dict[str, Any]) -> Tuple[str, str]:
    data = cv.get("profile", {})
    return data.get("name") or "Unnamed CV", data.get("role") or ""


def display_name(name: str) -> str:
    parts = (name or "").strip().split(maxsplit=1)
    if len(parts) == 2:
        return f"{parts[0]}\n{parts[1]}"
    return name or ""


def fit_profile_name_lines(
    first_name: str,
    last_name: str,
    *,
    max_width: float,
    measure: Callable[[str, float], float],
    base_size: float,
    min_size: float = 16.0,
    safety_factor: float = 0.96,
) -> Tuple[str, str, float]:
    """Fit a profile name into at most two explicit display lines.

    Keep the semantic first/last split when both lines fit. If either line is
    too wide, rebalance the complete name at a word boundary, then reduce only
    the name size when the best two-line split still needs it. Renderers can
    disable automatic wrapping after applying this result, making a third
    header line structurally impossible without changing the stored name.
    """
    first = re.sub(r"\s+", " ", str(first_name or "")).strip()
    last = re.sub(r"\s+", " ", str(last_name or "")).strip()
    original = tuple(part for part in (first, last) if part)
    if not original:
        return "", "", base_size

    safe_width = max(1.0, max_width * safety_factor)

    def widest(lines: Tuple[str, ...], size: float) -> float:
        return max(measure(line, size) for line in lines)

    if len(original) <= 2 and widest(original, base_size) <= safe_width:
        lines = original
    else:
        words = " ".join(original).split()
        if len(words) == 1:
            lines = (words[0],)
        else:
            candidates = [
                (" ".join(words[:index]), " ".join(words[index:]))
                for index in range(1, len(words))
            ]
            lines = min(
                candidates,
                key=lambda candidate: (
                    widest(candidate, base_size),
                    abs(
                        measure(candidate[0], base_size)
                        - measure(candidate[1], base_size)
                    ),
                ),
            )

    fitted_size = float(base_size)
    while fitted_size > min_size and widest(lines, fitted_size) > safe_width:
        fitted_size = max(min_size, fitted_size - 0.5)
    # ``min_size`` is the preferred design floor, not permission to create a
    # third line. An unusually long real-world name may go below it so the
    # renderer can still honour the hard two-line header contract.
    remaining_width = widest(lines, fitted_size)
    if remaining_width > safe_width:
        fitted_size *= safe_width / remaining_width * 0.98

    line_one = lines[0]
    line_two = lines[1] if len(lines) > 1 else ""
    return line_one, line_two, fitted_size


def overview(cv: Dict[str, Any], limit: int = 1300) -> str:
    return clamp(section_text(cv, "Overview"), limit)


def key_expertise(cv: Dict[str, Any], limit: int = 8) -> List[str]:
    bullets = section_bullets(cv, "Core Competencies")
    if not bullets:
        bullets = section_lines(cv, "Qualifications")[:limit]
    return [clean_display_line(item) for item in bullets[:limit] if clean_display_line(item)]


def qualifications(cv: Dict[str, Any], limit: int = 8) -> List[str]:
    lines = section_lines(cv, "Qualifications")
    education = section_lines(cv, "Education")
    combined = lines + education
    return [clean_display_line(item) for item in combined[:limit] if clean_display_line(item)]


def memberships(cv: Dict[str, Any], limit: int = 5) -> List[str]:
    lines = section_lines(
        cv,
        "Professional Bodies and Memberships",
        "Professional Bodies",
        "Professional Memberships",
        "Memberships",
    )[:limit]
    return [clean_display_line(item) for item in lines if clean_display_line(item)]


def work_experience(cv: Dict[str, Any], limit: int = 7) -> List[str]:
    section = get_section(cv, "Work Experience")
    if not section:
        return []
    entries = section.get("work_experience") or []
    if entries:
        formatted = []
        for entry in entries:
            years = entry.get("start_year") or ""
            if entry.get("end_year"):
                years = f"{years} - {entry['end_year']}"
            role = entry.get("role") or ""
            org = entry.get("organisation") or ""
            formatted.append(", ".join(part for part in [years, role, org] if part))
        return [clean_display_line(item) for item in formatted[:limit] if clean_display_line(item)]
    return [clean_display_line(item) for item in section_lines(cv, "Work Experience")[:limit] if clean_display_line(item)]


def selected_projects(cv: Dict[str, Any], limit: Optional[int] = 12) -> List[Dict[str, Any]]:
    section = get_section(cv, "Selected Projects", "Selected Project Highlights")
    if not section:
        return []
    projects = section.get("projects") or []
    projects = normalise_project_entries(projects)
    return projects[:limit] if limit else projects


def normalise_project_entries(projects: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    cleaned: List[Dict[str, Any]] = []
    for project in projects:
        name = re.sub(r"\s+", " ", (project.get("name") or "").strip())
        if not name:
            continue
        project = copy.deepcopy(project)
        project["name"] = name
        if is_project_table_header(name):
            continue
        if is_project_intro_name(name):
            continue
        if cleaned and is_project_fragment_name(name):
            append_project_fragment(cleaned[-1], project)
            continue
        cleaned.append(project)
    return cleaned


def is_project_table_header(name: str) -> bool:
    normalized = re.sub(r"[^A-Za-z]+", " ", name).strip().lower()
    return normalized in {
        "key projects location client certification",
        "projects location client certification",
        "key projects",
    }


def is_project_intro_name(name: str) -> bool:
    """Reject narrative lead-ins that extraction labelled as project names."""
    stripped = re.sub(r"\s+", " ", name).strip()
    if re.match(r"^The following projects?\b", stripped, re.I):
        return True
    if re.match(
        r"^[A-Z][A-Za-z.' -]{1,70}\s+has\s+"
        r"(?:overseen|delivered|worked on|completed|contributed to)\b",
        stripped,
        re.I,
    ) and re.search(r"\bprojects?\b", stripped, re.I):
        return True
    return False


def is_project_fragment_name(name: str) -> bool:
    stripped = name.strip()
    if not stripped:
        return True
    first = stripped[0]
    if first.islower():
        return True
    if first.isdigit() and re.match(r"^\d+\s+(?:targets|objectives|requirements|recommendations|actions)\b", stripped, re.I):
        return True
    if re.match(r"^Developed\b", stripped) and not re.search(r",\s*(?:19|20)\d{2}\s*[-–]", stripped):
        return True
    # A page break can promote the first words of a wrapped description into
    # the project-name column. These fragments are prose, not new projects.
    if re.match(r"^(?:19|20)\d{2}\s+[a-z]", stripped):
        return True
    if re.match(
        r"^[^.]{2,55}\.\s+(?:The|This|Its|His|Her|A|An)\b",
        stripped,
    ):
        return True
    if re.fullmatch(r"(?:[A-Za-z]+ )?\d{4}\s*[-–]\s*(?:[A-Za-z]+ )?\d{4}|(?:[A-Za-z]+ )?\d{4}\s*[-–]\s*ongoing", stripped, re.I):
        return True
    if re.fullmatch(r"[A-Za-z]+ \d{4}\s*[-–]\s*[A-Za-z]+ \d{4}", stripped):
        return True
    if re.fullmatch(r"[A-Za-z]+ \d{4}", stripped):
        return True
    if len(stripped) > 115 and not re.search(r"\b(?:19|20)\d{2}\b", stripped):
        return True
    words = stripped.split()
    if len(words) <= 2 and stripped.upper() == stripped and not re.search(r"\d", stripped):
        project_words = {
            "STADIUM",
            "STADIUMS",
            "BRIDGE",
            "BRIDGES",
            "HOTEL",
            "PALACE",
            "VILLAS",
            "TOWER",
            "TOWERS",
            "MALL",
            "PARK",
            "DISTRICT",
            "SEDRA",
        }
        if not any(word.strip("&|") in project_words for word in words):
            return True
    if re.match(r"^(?:PIF|KSA|UAE|UK|BSBG|DGDA|RICS|JT|Parsons|Partners)(?:\b| \|)", stripped):
        return True
    return False


def append_project_fragment(previous: Dict[str, Any], fragment: Dict[str, Any]) -> None:
    fragment_parts = [fragment.get("name", "")]
    fragment_parts.extend(fragment.get("bullets") or [])
    if fragment.get("text"):
        fragment_parts.append(fragment["text"])
    addition = " ".join(part for part in fragment_parts if part).strip()
    if not addition:
        return
    bullets = previous.setdefault("bullets", [])
    if bullets:
        bullets[-1] = f"{bullets[-1]} {addition}".strip()
    else:
        bullets.append(addition)


def other_sections(cv: Dict[str, Any], limit: int = 3) -> List[Dict[str, Any]]:
    sections = []
    for section in cv.get("sections", []):
        key = section.get("heading", "").lower()
        if key in STANDARD_SECTION_KEYS:
            continue
        if key in {
            "strategic & frameworks",
            "hospitality & leisure",
            "masterplans & communities",
            "mixed-use",
            "sports & stadiums",
            "commercial",
            "residential",
            "cultural & heritage",
            "infrastructure & utilities",
        }:
            continue
        sections.append(section)
    return sections[:limit]


def project_title(project: Dict[str, Any], limit: int = 145) -> str:
    return clamp(project.get("name", "Project"), limit)


def project_body(project: Dict[str, Any], limit: int = 450) -> str:
    bullets = project.get("bullets") or []
    body = " ".join(bullets) if bullets else project.get("text", "")
    return clamp(body, limit)


def project_body_full(project: Dict[str, Any]) -> str:
    bullets = project.get("bullets") or []
    body = " ".join(bullets) if bullets else project.get("text", "")
    return re.sub(r"\s+", " ", (body or "")).strip()


def project_body_limit(project_count: int, base: int, floor: int = 140) -> int:
    if project_count <= 6:
        return base
    if project_count <= 12:
        return max(floor, int(base * 0.78))
    if project_count <= 24:
        return max(floor, int(base * 0.58))
    return floor


def chunks(items: List[Any], size: int) -> List[List[Any]]:
    return [items[index : index + size] for index in range(0, len(items), size)]


def estimate_wrapped_lines(text: str, width_emu: int, font_size_pt: float) -> int:
    """Estimate wrapped lines for Office text where the real layout engine is unavailable."""
    if not text:
        return 1
    width_inches = max(0.5, width_emu / 914400)
    chars_per_line = max(18, int((width_inches * 92) / max(font_size_pt, 1)))
    lines = 0
    for raw_line in str(text).splitlines() or [""]:
        line = raw_line.strip()
        lines += max(1, math.ceil(len(line) / chars_per_line))
    return lines


def estimate_pptx_text_height_emu(
    title: str,
    body: str,
    width_emu: int,
    title_size_pt: float,
    body_size_pt: float,
    padding_pt: float = 12,
) -> int:
    title_lines = estimate_wrapped_lines(title, width_emu, title_size_pt)
    body_lines = estimate_wrapped_lines(body, width_emu, body_size_pt)
    height_pt = padding_pt + (title_lines * title_size_pt * 1.22) + (body_lines * body_size_pt * 1.16)
    return int(height_pt * 12700)


def copy_docx_paragraph_format(target: Any, prototype: Any) -> None:
    if prototype is None:
        return
    try:
        target.style = prototype.style
    except Exception:
        pass
    p_pr = prototype._p.pPr
    if p_pr is not None:
        existing = target._p.pPr
        if existing is not None:
            target._p.remove(existing)
        target._p.insert(0, copy.deepcopy(p_pr))


def copy_docx_run_format(target: Any, prototype: Any) -> None:
    if prototype is None:
        return
    r_pr = prototype._r.rPr
    if r_pr is not None:
        existing = target._r.rPr
        if existing is not None:
            target._r.remove(existing)
        target._r.insert(0, copy.deepcopy(r_pr))


def clear_docx_cell(cell: Any) -> None:
    cell._tc.clear_content()


def remove_docx_table(table: Any) -> None:
    parent = table._element.getparent()
    if parent is not None:
        parent.remove(table._element)


def remove_docx_table_row(row: Any) -> None:
    parent = row._tr.getparent()
    if parent is not None:
        parent.remove(row._tr)


def set_docx_cell_text_from_prototype(cell: Any, text: str, prototype: Optional[Any] = None) -> Any:
    clear_docx_cell(cell)
    paragraph = cell.add_paragraph()
    if prototype is not None:
        copy_docx_paragraph_format(paragraph, prototype)
    run = paragraph.add_run(text)
    if prototype is not None and prototype.runs:
        copy_docx_run_format(run, prototype.runs[0])
    return paragraph


def set_docx_cell_border(cell: Any, **borders: Dict[str, str]) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        attrs = borders.get(edge)
        if attrs is None:
            continue
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in attrs.items():
            element.set(qn(f"w:{key}"), str(value))


def set_docx_table_borders_none(table: Any) -> None:
    border = {"val": "nil"}
    for row in table.rows:
        for cell in row.cells:
            set_docx_cell_border(
                cell,
                top=border,
                left=border,
                bottom=border,
                right=border,
                insideH=border,
                insideV=border,
            )


def convert_docx_text_newlines(doc: Any) -> None:
    for text_node in list(doc.element.xpath(".//w:t")):
        text = text_node.text
        if not text or "\n" not in text:
            continue
        lines = text.splitlines()
        text_node.text = lines[0] if lines else ""
        parent_run = text_node.getparent()
        if parent_run is None:
            continue
        for line in lines[1:]:
            parent_run.append(OxmlElement("w:br"))
            extra_text = OxmlElement("w:t")
            extra_text.text = line
            parent_run.append(extra_text)


def duplicate_pptx_slide(prs: Any, slide_index: int) -> Any:
    source = prs.slides[slide_index]
    blank_layout = prs.slide_layouts[len(prs.slide_layouts) - 1]
    destination = prs.slides.add_slide(blank_layout)
    for shape in source.shapes:
        destination.shapes._spTree.insert_element_before(
            copy.deepcopy(shape.element),
            "p:extLst",
        )
    return destination


def delete_pptx_slide(prs: Any, slide_index: int) -> None:
    slide_id_list = prs.slides._sldIdLst
    slide_ids = list(slide_id_list)
    rel_id = slide_ids[slide_index].rId
    prs.part.drop_rel(rel_id)
    slide_id_list.remove(slide_ids[slide_index])


def clamp(text: str, limit: int) -> str:
    text = re.sub(r"\s+", " ", (text or "")).strip()
    if len(text) <= limit:
        return text
    return text[: max(0, limit - 1)].rstrip() + "…"


def cv_photo_bytes(cv: Dict[str, Any]) -> Optional[bytes]:
    photo = cv.get("profile_photo") or {}
    encoded = photo.get("base64")
    if not encoded:
        return None
    try:
        return base64.b64decode(encoded)
    except Exception:
        return None


def write_temp_photo(cv: Dict[str, Any], directory: Path, square: bool = False) -> Optional[Path]:
    data = cv_photo_bytes(cv)
    if not data:
        return None
    suffix = ".png" if square else ".jpg"
    path = directory / f"{slugify(profile(cv)[0])}{suffix}"
    if square:
        source = directory / f"{slugify(profile(cv)[0])}_source"
        source.write_bytes(data)
        image = Image.open(source).convert("RGB")
        image = ImageOps.fit(image, (700, 700), method=Image.Resampling.LANCZOS)
        image.save(path, format="PNG")
    else:
        path.write_bytes(data)
    return path


def patch_docx_profile_photo(docx_path: Path, cv: Dict[str, Any]) -> None:
    data = cv_photo_bytes(cv)
    if not data:
        return
    with tempfile.TemporaryDirectory() as tmp:
        tmpdir = Path(tmp)
        source = tmpdir / "source"
        source.write_bytes(data)
        image = Image.open(source).convert("RGB")
        image = ImageOps.fit(image, (700, 700), method=Image.Resampling.LANCZOS)
        output = tmpdir / "image1.png"
        image.save(output, format="PNG")
        replacement = output.read_bytes()

    original = docx_path.read_bytes()
    temp_path = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(temp_path, "w") as zout:
        for item in zin.infolist():
            content = zin.read(item.filename)
            if item.filename == "word/media/image1.png":
                content = replacement
            zout.writestr(item, content)
    temp_path.replace(docx_path)


def patch_docx_name_title_box_height(docx_path: Path, min_height_emu: int = 1700000) -> None:
    with zipfile.ZipFile(docx_path, "r") as archive:
        original_items = [(item, archive.read(item.filename)) for item in archive.infolist()]

    next_doc_pr_id = 12000
    namespaces = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "v": "urn:schemas-microsoft-com:vml",
    }
    w_val = f"{{{namespaces['w']}}}val"

    def run_has_style(run: Any, style_id: str) -> bool:
        return bool(run.xpath(f".//w:pStyle[@w:val='{style_id}']", namespaces=namespaces))

    def remove_paragraphs_by_style(root: Any, style_id: str) -> None:
        for paragraph in list(root.xpath(f".//w:p[w:pPr/w:pStyle[@w:val='{style_id}']]", namespaces=namespaces)):
            parent = paragraph.getparent()
            if parent is not None:
                parent.remove(paragraph)

    def remove_empty_name_paragraphs(root: Any) -> None:
        for paragraph in list(root.xpath(".//w:p[w:pPr/w:pStyle[@w:val='AESGNameTitle']]", namespaces=namespaces)):
            text = "".join(paragraph.xpath(".//w:t/text()", namespaces=namespaces)).strip()
            if not text:
                parent = paragraph.getparent()
                if parent is not None:
                    parent.remove(paragraph)

    def set_box_height(root: Any, height_emu: int, exact: bool = False) -> None:
        for element in root.xpath(".//wp:extent | .//a:ext", namespaces=namespaces):
            current = int(element.get("cy") or 0)
            if exact or current < height_emu:
                element.set("cy", str(height_emu))

    def move_anchor_down(root: Any, offset_emu: int) -> None:
        offsets = root.xpath(".//wp:positionV/wp:posOffset", namespaces=namespaces)
        if offsets:
            offsets[0].text = str(int(offsets[0].text or "0") + offset_emu)

    def refresh_doc_pr(root: Any) -> None:
        nonlocal next_doc_pr_id
        next_doc_pr_id += 1
        for doc_pr in root.xpath(".//wp:docPr", namespaces=namespaces):
            doc_pr.set("id", str(next_doc_pr_id))
            doc_pr.set("name", f"Role Text Box {next_doc_pr_id}")
            break
        for shape in root.xpath(".//v:shape", namespaces=namespaces):
            shape.set("id", f"_x0000_s{next_doc_pr_id}")
            break

    def patch_document_xml(content: bytes) -> bytes:
        root = etree.fromstring(content)
        runs = list(root.xpath(".//w:r[w:drawing]", namespaces=namespaces))
        for run in runs:
            has_name = run_has_style(run, "AESGNameTitle")
            has_role = run_has_style(run, "AESGCandidateTitle")
            if has_name and has_role:
                parent = run.getparent()
                if parent is None:
                    continue
                name_run = copy.deepcopy(run)
                role_run = copy.deepcopy(run)

                remove_paragraphs_by_style(name_run, "AESGCandidateTitle")
                remove_empty_name_paragraphs(name_run)
                set_box_height(name_run, min_height_emu)

                remove_paragraphs_by_style(role_run, "AESGNameTitle")
                set_box_height(role_run, 560000, exact=True)
                move_anchor_down(role_run, 950000)
                refresh_doc_pr(role_run)

                index = parent.index(run)
                parent.remove(run)
                parent.insert(index, name_run)
                parent.insert(index + 1, role_run)
            elif has_name or has_role:
                set_box_height(run, min_height_emu)
        return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)

    temp_path = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(temp_path, "w") as archive:
        for item, content in original_items:
            if item.filename == "word/document.xml":
                content = patch_document_xml(content)
            archive.writestr(item, content)
    temp_path.replace(docx_path)


def set_docx_styled_text(doc: Any, style_ids: Iterable[str], text: str) -> bool:
    changed = False
    wanted = set(style_ids)
    for paragraph in doc.element.xpath(".//w:p"):
        p_pr = paragraph.find(qn("w:pPr"))
        if p_pr is None:
            continue
        p_style = p_pr.find(qn("w:pStyle"))
        if p_style is None or p_style.get(qn("w:val")) not in wanted:
            continue
        text_nodes = paragraph.xpath(".//w:t")
        if not text_nodes:
            continue
        lines = str(text).splitlines() or [""]
        text_nodes[0].text = lines[0]
        parent_run = text_nodes[0].getparent()
        for line in lines[1:]:
            parent_run.append(OxmlElement("w:br"))
            extra_text = OxmlElement("w:t")
            extra_text.text = line
            parent_run.append(extra_text)
        for node in text_nodes[1:]:
            node.text = ""
        changed = True
    return changed


def remove_docx_placeholder_paragraphs(doc: Any) -> None:
    markers = (
        "Project name",
        "Insert Candidate Role",
        "Pereperi",
        "Role:",
    )
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if any(marker in text for marker in markers):
            parent = paragraph._element.getparent()
            if parent is not None:
                parent.remove(paragraph._element)


def add_parser(
    description: str,
    default_template: str,
    default_output_dir: str,
    default_map: Optional[str] = None,
) -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=description)
    parser.add_argument("--data", type=Path, default=Path("output/cv_data.json"))
    parser.add_argument("--template", type=Path, default=Path(default_template))
    parser.add_argument("--template-map", type=Path, default=Path(default_map) if default_map else None)
    parser.add_argument("--output-dir", type=Path, default=Path(default_output_dir))
    parser.add_argument("--limit", type=int, default=None, help="Optional maximum number of CVs to generate.")
    return parser


def iter_cvs(data_path: Path, limit: Optional[int] = None) -> List[Dict[str, Any]]:
    cvs = load_cvs(data_path)
    return cvs[:limit] if limit else cvs
